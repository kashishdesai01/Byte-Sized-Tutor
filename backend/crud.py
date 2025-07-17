from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session
import shutil, os, random
from typing import Optional, List
from pydantic import BaseModel, Field
from sqlalchemy import func


from . import models, schemas, auth
from langchain_core.messages import HumanMessage, AIMessage

import fitz
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAI
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder, PromptTemplate
from langchain.chains.summarize import load_summarize_chain
from langchain_core.output_parsers import JsonOutputParser

# --- Model & Directory Initialization ---
UPLOAD_DIRECTORY = "./uploads"
VECTOR_STORE_DIRECTORY = "./vector_stores"
embedding_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
llm = GoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.7)

# --- Helper functions ---
def get_user_from_db(db: Session, email: str):
    return db.query(models.User).filter(models.User.email == email).first()

def get_document_from_db(db: Session, doc_id: int):
    document = db.query(models.Document).filter(models.Document.id == doc_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found.")
    return document


async def create_document(db: Session, file: UploadFile, user: Optional[dict]):
    owner_id = None
    if user:
        db_user = get_user_from_db(db, user["email"])
        if db_user:
            owner_id = db_user.id

    temp_file_path = os.path.join(UPLOAD_DIRECTORY, file.filename)
    try:
        with open(temp_file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        full_text = ""
        if file.content_type == "application/pdf":
            doc = fitz.open(temp_file_path)
            full_text = "".join(page.get_text() for page in doc)
        else:
            full_text = "\n".join([p.text for p in docx.Document(temp_file_path).paragraphs])

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(full_text)
        if not chunks:
            raise ValueError("Document could not be chunked.")
            
        vector_store = FAISS.from_texts(chunks, embedding_model)

        db_document = models.Document(filename=file.filename, owner_id=owner_id)
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        
        vector_store_path = os.path.join(VECTOR_STORE_DIRECTORY, f"doc_{db_document.id}")
        vector_store.save_local(vector_store_path)

        return db_document
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)

def get_user_documents(db: Session, user_email: str):
    user = get_user_from_db(db, user_email)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return db.query(models.Document).filter(models.Document.owner_id == user.id).all()

def get_chat_history(db: Session, document_id: int):
    return db.query(models.ChatHistory).filter(models.ChatHistory.document_id == document_id).order_by(models.ChatHistory.timestamp).all()

def create_chat_message(db: Session, document_id: int, role: str, content: str, user_id: Optional[int] = None):
    db_message = models.ChatHistory(
        document_id=document_id,
        user_id=user_id,
        role=role,
        content=content
    )
    db.add(db_message)
    db.commit()
    db.refresh(db_message)
    return db_message

async def get_answer(db: Session, request: schemas.AskRequest, user: Optional[dict]):
    document = get_document_from_db(db, request.document_id)
    vector_store_path = os.path.join(VECTOR_STORE_DIRECTORY, f"doc_{document.id}")
    if not os.path.exists(vector_store_path):
        raise HTTPException(status_code=404, detail="Vector store not found.")
    
    vector_store = FAISS.load_local(
        vector_store_path, embedding_model, allow_dangerous_deserialization=True
    )
    retriever = vector_store.as_retriever()

    chat_history_messages = [
        HumanMessage(content=msg.content) if msg.role == "human" else AIMessage(content=msg.content)
        for msg in request.chat_history or []
    ]

    base_persona = """
    You are the AI Study Buddy, an expert tutor. Your primary goal is to help a user understand the provided context by explaining it clearly and conversationally.

**CRITICAL FORMATTING RULES:**
    - Your entire response MUST use Markdown for all formatting.
    - Use headings (`#`, `##`), bold (`**text**`), italics (`*text*`), and lists (`-`, `1.`) to structure your explanation logically.
    - When explaining a formula, you MUST place the formula on its own line in a Markdown code block. On the lines immediately following, use a bulleted list to define what each symbol in the formula represents. For example:
    
    The formula for kinetic energy is:
    ```
    K = 1/2 * m * v^2
    ```
    Where:
    - `K` is the kinetic energy.
    - `m` is the mass of the object.
    - `v` is the velocity of the object.
    
    - For sub-topics or nested ideas, use indented bullet points to show the hierarchy of information.

    **TUTORING STYLE:**
    You are an intelligent, helpful, and highly knowledgeable assistant.
    You explain complex concepts clearly and thoroughly in a way that’s easy for beginners to understand.
    You adapt the level of detail based on how the user asks their question.
    All explanations MUST be based *only* on the provided context below. If the answer is not in the context, clearly state that.
    Always answer honestly and clearly. Use examples, analogies, or step-by-step reasoning when helpful.
    If applicable, explain any relevant formulas or technical terms mentioned in the text.
    
"""
    user_question_lower = request.question.lower()
    
    if any(kw in user_question_lower for kw in ["in detail", "detailed", "elaborate"]):
        detail_instruction = "Give a long, detailed explanation..."
    elif any(kw in user_question_lower for kw in ["in depth", "explain", "describe"]):
        detail_instruction = "Provide a thorough, multi-paragraph explanation..."
    else:
        detail_instruction = "Keep the answer concise..."

    qa_system_prompt = f"{base_persona}\n\n{detail_instruction}\n\n{{context}}"

    contextualize_q_prompt = ChatPromptTemplate.from_messages([
        ("system", "Given a chat history... reformulate it if needed..."),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}")
    ])
    history_aware_retriever = create_history_aware_retriever(llm, retriever, contextualize_q_prompt)

    qa_prompt = ChatPromptTemplate.from_messages([
        ("system", qa_system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}")
    ])
    question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)
    rag_chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)

    response = await rag_chain.ainvoke({
        "chat_history": chat_history_messages,
        "input": request.question
    })

    user_id = None
    if user:
        db_user = get_user_from_db(db, user["email"])
        if db_user:
            user_id = db_user.id

    create_chat_message(db, document.id, "human", request.question, user_id)
    create_chat_message(db, document.id, "ai", response["answer"], user_id)

    return {"answer": response["answer"]}



async def get_summary(db: Session, request: schemas.DocumentRequest):
    document = get_document_from_db(db, request.document_id)
    vector_store_path = os.path.join(VECTOR_STORE_DIRECTORY, f"doc_{document.id}")
    if not os.path.exists(vector_store_path):
        raise HTTPException(status_code=404, detail="Vector store not found.")

    vector_store = FAISS.load_local(
        vector_store_path, embedding_model, allow_dangerous_deserialization=True
    )
    docs = vector_store.similarity_search("", k=50)
    if not docs:
        raise HTTPException(status_code=404, detail="No content to summarize.")

    full_text = " ".join([doc.page_content for doc in docs])

    summary_prompt = PromptTemplate(
        input_variables=["text"],
        template="""
You are an intelligent and friendly assistant. Your task is to summarize the following content clearly and concisely, as if explaining it to someone who wants to quickly understand the main points.

Write a 3-paragraph summary that captures the key ideas. If the content is technical, simplify it a little for better readability. Avoid sounding robotic.

TEXT:
{text}

SUMMARY:
"""
    )

    chain = summary_prompt | llm
    summary = await chain.ainvoke({"text": full_text})

    return {"summary": summary}



async def create_quiz(db: Session, request: schemas.DocumentRequest):
    document = get_document_from_db(db, request.document_id)
    vector_store_path = os.path.join(VECTOR_STORE_DIRECTORY, f"doc_{document.id}")
    if not os.path.exists(vector_store_path): raise HTTPException(status_code=404, detail="Vector store not found.")

    vector_store = FAISS.load_local(vector_store_path, embedding_model, allow_dangerous_deserialization=True)
    all_docs = vector_store.similarity_search("", k=100)
    if not all_docs: raise HTTPException(status_code=404, detail="No content to create quiz from.")
    
    random.shuffle(all_docs)
    sample_docs = all_docs[:20]
    full_context = " ".join([doc.page_content for doc in sample_docs])

    parser = JsonOutputParser(pydantic_object=schemas.Quiz)
    
    prompt = PromptTemplate(
    template="""
You are a strict and intelligent quiz generation assistant. Your job is to create a high-quality, diverse 5-question multiple-choice quiz from the provided Source Text.

💡 Your questions can be about *anything* — programming, biology, physics, law, history, etc. — depending on the content.

⚠️ **IMPORTANT RULES (MUST FOLLOW):**
1. **NO OUTSIDE KNOWLEDGE:** Use ONLY what’s in the Source Text below. Do not invent code, formulas, scenarios, or facts that are not present.
2. **STRICT GROUNDING:** Every question, all answer choices, and explanations must be directly answerable from the text. Do not reference “the provided code” unless actual code appears in the Source Text.
3. **DIVERSITY OF QUESTIONS:** Include a variety of question types: definitions, cause-effect, reasoning, comparisons, structure, logic, etc.
4. **ANSWERABLE QUESTIONS ONLY:** Each question should have one correct answer that is clearly supported by the Source Text.
5. **VALID FORMAT:** Return a JSON object in this exact schema: {format_instructions}. `correct_answer` must exactly match one of the `options`.

📄 **SOURCE TEXT**
---
{context}
---

Now generate the quiz.
""",
    input_variables=["context"],
    partial_variables={"format_instructions": parser.get_format_instructions()},
)

    chain = prompt | llm | parser
    quiz_data = await chain.ainvoke({"context": full_context})
    return quiz_data

def save_quiz_attempt(db: Session, user_id: int, request: schemas.SubmitQuizRequest):
   
    db_attempt = models.QuizAttempt(
        document_id=request.document_id,
        user_id=user_id,
        score=request.score
    )
    db.add(db_attempt)
    db.commit()
    db.refresh(db_attempt)

    
    for answer in request.answers:
        db_answer = models.QuizAnswer(
            attempt_id=db_attempt.id,
            question_text=answer.question_text,
            selected_answer=answer.selected_answer,
            correct_answer=answer.correct_answer,
            is_correct=answer.is_correct
        )
        db.add(db_answer)
    
    db.commit()
    return db_attempt

def get_quiz_history_for_document(db: Session, user_id: int, document_id: int):
    attempts = db.query(models.QuizAttempt).filter(
        models.QuizAttempt.user_id == user_id,
        models.QuizAttempt.document_id == document_id
    ).order_by(models.QuizAttempt.timestamp.desc()).all()
    return attempts


def delete_document(db: Session, document_id: int, user_id: int):
    document = db.query(models.Document).filter(models.Document.id == document_id, models.Document.owner_id == user_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found or access denied")
    
    # Delete the associated vector store file
    vector_store_path = os.path.join(VECTOR_STORE_DIRECTORY, f"doc_{document.id}")
    if os.path.exists(vector_store_path):
        shutil.rmtree(vector_store_path) # Use rmtree to delete the folder

    db.delete(document)
    db.commit()
    return {"message": "Document and all associated data deleted successfully."}

def delete_chat_history(db: Session, document_id: int, user_id: int):
    document = db.query(models.Document).filter(models.Document.id == document_id, models.Document.owner_id == user_id).first()
    if not document:
        raise HTTPException(status_code=403, detail="Not authorized to delete this chat history")
    
    db.query(models.ChatHistory).filter(models.ChatHistory.document_id == document_id).delete(synchronize_session=False)
    db.commit()
    return {"message": "Chat history deleted successfully."}

def delete_all_quiz_history(db: Session, document_id: int, user_id: int):
    document = db.query(models.Document).filter(models.Document.id == document_id, models.Document.owner_id == user_id).first()
    if not document:
        raise HTTPException(status_code=403, detail="Not authorized to delete this quiz history")
    
    attempts_to_delete = db.query(models.QuizAttempt).filter(models.QuizAttempt.document_id == document_id, models.QuizAttempt.user_id == user_id).all()
    
    for attempt in attempts_to_delete:
        db.delete(attempt)
        
    db.commit()
    return {"message": "All quiz history for this document has been deleted."}

def delete_single_quiz_attempt(db: Session, attempt_id: int, user_id: int):
    attempt = db.query(models.QuizAttempt).filter(models.QuizAttempt.id == attempt_id, models.QuizAttempt.user_id == user_id).first()
    if not attempt:
        raise HTTPException(status_code=404, detail="Quiz attempt not found or access denied")
    
    db.delete(attempt)
    db.commit()
    return {"message": "Quiz attempt deleted successfully."}

def delete_multiple_quiz_attempts(db: Session, attempt_ids: List[int], user_id: int):
    # Fetch attempts to ensure user has permission for all of them
    attempts_to_delete = db.query(models.QuizAttempt).filter(
        models.QuizAttempt.id.in_(attempt_ids),
        models.QuizAttempt.user_id == user_id
    ).all()

    if len(attempts_to_delete) != len(attempt_ids):
        raise HTTPException(status_code=403, detail="One or more quiz attempts not found or access denied")

    for attempt in attempts_to_delete:
        db.delete(attempt)

    db.commit()
    return {"message": f"{len(attempts_to_delete)} quiz attempts deleted successfully."}

def get_flashcard_sets_for_document(db: Session, user_id: int, document_id: int):
    return db.query(models.FlashcardSet).filter(
        models.FlashcardSet.user_id == user_id,
        models.FlashcardSet.document_id == document_id
    ).order_by(models.FlashcardSet.timestamp.desc()).all()

async def create_flashcards(db: Session, document_id: int, user_id: int):
    document = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found.")

    vector_store_path = os.path.join(VECTOR_STORE_DIRECTORY, f"doc_{document.id}")
    if not os.path.exists(vector_store_path):
        raise HTTPException(status_code=404, detail="Vector store for document not found.")

    vector_store = FAISS.load_local(vector_store_path, embedding_model, allow_dangerous_deserialization=True)
    all_docs = vector_store.similarity_search("", k=100)
    if not all_docs:
        raise HTTPException(status_code=404, detail="No content found to create flashcards from.")
        
    full_context = " ".join([doc.page_content for doc in all_docs])

    # Define a Pydantic model for the AI to structure its output
    class Flashcard(BaseModel):
        term: str = Field(description="The key term, concept, or name.")
        definition: str = Field(description="A clear and concise definition or explanation of the term.")

    class FlashcardSet(BaseModel):
        flashcards: List[Flashcard] = Field(description="A list of generated flashcards.")

    parser = JsonOutputParser(pydantic_object=FlashcardSet)

    prompt = PromptTemplate(
        template="""You are an expert educator specializing in creating effective study materials.
        Based *only* on the provided text, identify the most important key terms, concepts, and definitions.
        Generate a set of 10 high-quality flashcards from this text.
        Your output must be a JSON object that strictly follows this format: {format_instructions}

        **Source Text:**
        ---
        {context}
        ---
        """,
        input_variables=["context"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )

    chain = prompt | llm | parser
    generated_data = await chain.ainvoke({"context": full_context})
    
    # Save the new flashcard set to the database
    new_set = models.FlashcardSet(
        document_id=document_id,
        user_id=user_id,
        title=f"Flashcards for {document.filename}"
    )
    db.add(new_set)
    db.commit()
    db.refresh(new_set)

    for card_data in generated_data['flashcards']:
        db_card = models.Flashcard(
            set_id=new_set.id,
            front=card_data['term'],
            back=card_data['definition']
        )
        db.add(db_card)
    
    db.commit()
    return new_set


def delete_flashcard_set(db: Session, set_id: int, user_id: int):
    """Deletes a single flashcard set after verifying ownership."""
    set_to_delete = db.query(models.FlashcardSet).filter(models.FlashcardSet.id == set_id).first()

    if not set_to_delete:
        raise HTTPException(status_code=404, detail="Flashcard set not found")

    if set_to_delete.user_id != user_id:
        raise HTTPException(status_code=403, detail="Not authorized to delete this flashcard set")

    db.delete(set_to_delete)
    db.commit()
    return

def delete_multiple_flashcard_sets(db: Session, item_ids: List[int], user_id: int):
    """Deletes multiple flashcard sets by their IDs after verifying ownership."""
    sets_to_delete = db.query(models.FlashcardSet).filter(
        models.FlashcardSet.id.in_(item_ids),
        models.FlashcardSet.user_id == user_id
    ).all()

    if len(sets_to_delete) != len(set(item_ids)):
         raise HTTPException(status_code=403, detail="One or more flashcard sets not found or access denied")

    for f_set in sets_to_delete:
        db.delete(f_set)

    db.commit()
    return

def delete_all_flashcard_sets_for_document(db: Session, document_id: int, user_id: int):
    """Deletes all flashcard sets for a document after verifying ownership."""
    document = db.query(models.Document).filter(
        models.Document.id == document_id,
        models.Document.owner_id == user_id
    ).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found or you do not have permission")

    document.flashcard_sets = []
    db.commit()
    return

def update_user_password(db: Session, user: models.User, new_password: str):
    """Updates the password for a given user object."""
    hashed_password = auth.get_password_hash(new_password)
    user.hashed_password = hashed_password
    db.commit()
    db.refresh(user)
    return user

def get_progress_report_for_document(db: Session, user_id: int, document_id: int):
    attempts = db.query(models.QuizAttempt).filter(
        models.QuizAttempt.user_id == user_id,
        models.QuizAttempt.document_id == document_id
    ).order_by(models.QuizAttempt.timestamp.asc()).all()

    if not attempts:
        return {
            "total_quizzes_taken": 0,
            "average_score": 0,
            "highest_score": 0,
            "scores_over_time": []
        }

    total_quizzes_taken = len(attempts)
    
    stats = db.query(
        func.avg(models.QuizAttempt.score).label('average_score'),
        func.max(models.QuizAttempt.score).label('highest_score')
    ).filter(
        models.QuizAttempt.user_id == user_id,
        models.QuizAttempt.document_id == document_id
    ).one()

    scores_over_time = [{"timestamp": attempt.timestamp, "score": attempt.score} for attempt in attempts]

    return {
        "total_quizzes_taken": total_quizzes_taken,
        "average_score": stats.average_score,
        "highest_score": stats.highest_score,
        "scores_over_time": scores_over_time
    }

